# -*- coding: utf-8 -*-
"""Сборка DZ-9 docx: руководство/ПЗ и отчёт для сдачи."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

ROOT = Path(__file__).resolve().parent


def setup_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def meta_lines(doc: Document, rows: list[tuple[str, str]]) -> None:
    for label, val in rows:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(val)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build() -> Path:
    doc = setup_doc("ДЗ-9 — руководство по эксплуатации и пояснительная записка")
    meta_lines(
        doc,
        [
            ("Студент:", "Игорь Кашинцев"),
            ("Курс:", "VibeCoder, Тема 9 — локальная CRM, Supabase, RLS"),
            ("Проект:", "DesignStudio CRM"),
            ("GitHub:", "https://github.com/treshkash323-alt/designstudio-crm"),
            ("Папка:", "Projects/ДЗ-9/designstudio-crm/"),
            ("Дата:", "08.06.2026"),
            ("Полная версия (MD):", "DZ-9_РУКОВОДСТВО_И_ПЗ.md"),
        ],
    )

    doc.add_heading("Аннотация", level=1)
    doc.add_paragraph(
        "Пояснительная записка и руководство по локальной CRM: Docker, Supabase, RLS, "
        "роли manager/admin, безопасность и чеклист сдачи. Код на GitHub без секретов."
    )

    doc.add_heading("1. Что это за проект", level=1)
    doc.add_paragraph(
        "Локальная CRM DesignStudio: менеджер видит только своих лидов, админ — всех. "
        "Стек: Docker + Supabase CLI + Next.js 14 + PostgreSQL RLS."
    )

    doc.add_heading("2. Глоссарий (кратко)", level=1)
    add_table(
        doc,
        ["Термин", "Простыми словами"],
        [
            ("Docker", "Контейнеры Supabase на вашем ПК"),
            ("Supabase CLI", "npx supabase start — БД и Auth локально"),
            ("Studio", "http://localhost:54323 — таблицы, SQL, пользователи"),
            ("RLS", "Правила PostgreSQL: кто какие строки видит"),
            ("JWT", "Токен после входа; role admin в app_metadata"),
            ("anon key", "Публичный ключ в .env.local (с RLS)"),
            ("service_role", "Секретный ключ только на сервере"),
        ],
    )

    doc.add_heading("3. Ежедневный запуск", level=1)
    add_numbered(
        doc,
        [
            "Docker Desktop → Engine running",
            "cd Projects\\ДЗ-9\\designstudio-crm",
            "npx supabase start (если ещё не запущен)",
            "npm run dev → http://localhost:3000",
        ],
    )
    doc.add_paragraph("Остановка: npx supabase stop")

    doc.add_heading("4. Роли и проверка", level=1)
    add_table(
        doc,
        ["Пользователь", "Страница", "Ожидание"],
        [
            ("manager@designstudio.ru", "/dashboard", "2 лида"),
            ("admin@designstudio.ru", "/admin/clients", "5 лидов"),
        ],
    )
    doc.add_paragraph(
        "После назначения admin через SQL — выйти и войти снова (обновить JWT)."
    )

    doc.add_heading("5. SQL после регистрации", level=1)
    doc.add_paragraph(
        "Файл supabase/seed_after_signup.sql — в Studio → SQL Editor → Run. "
        "Не запускать путь к файлу в PowerShell как команду."
    )

    doc.add_heading("6. Безопасность (кратко)", level=1)
    add_table(
        doc,
        ["Проверка", "Результат"],
        [
            (".env.local в git", "Нет — только .env.local.example"),
            ("RLS leads/profiles", "Да — manager 2 / admin 5"),
            ("Middleware /admin", "Да — manager редирект"),
            ("service_role в браузере", "Нет — admin.ts не в UI"),
            ("Security Advisor", "0 errors, 2 warnings (handle_new_user)"),
            ("Доступ из интернета", "Нет — localhost"),
        ],
    )
    doc.add_paragraph("Подробно: SECURITY.md в репозитории.")

    doc.add_heading("7. Сдача ДЗ — чеклист", level=1)
    doc.add_paragraph("Обязательные скрины:", style="List Bullet")
    add_bullets(
        doc,
        [
            "Менеджер: /dashboard, 2 клиента, RLS «(2)»",
            "Админ: /admin/clients, 5 строк, бейдж АДМИН",
            "Желательно: Docker, Studio → leads, Advisor 0 errors",
            "На скринах замазать anon/service keys при необходимости",
        ],
    )
    doc.add_paragraph("Документы: DZ-9_отчёт_для_сдачи.docx, SECURITY.md, TODO_ROADMAP.md")

    doc.add_heading("8. Развитие после ДЗ (TODO)", level=1)
    add_bullets(
        doc,
        [
            "EDU (AIKIVAVIORA): student / teacher / admin, курсы, RLS — см. TODO_ROADMAP.md",
            "Tilda: лендинг и интеграция — отдельное обсуждение, не блокирует ДЗ-9",
            "Prod: invite-only, уникальные ключи, CI — см. SECURITY.md §4",
        ],
    )

    doc.add_heading("9. Заключение", level=1)
    doc.add_paragraph(
        "ДЗ выполнено: локальный Supabase, RLS, два пользователя, разграничение 2 vs 5 лидов. "
        "Код на GitHub без секретов. Подробности — в DZ-9_РУКОВОДСТВО_И_ПЗ.md."
    )

    out = ROOT / "DZ-9_РУКОВОДСТВО_И_ПЗ.docx"
    doc.save(out)
    return out


def build_report() -> Path:
    doc = setup_doc("ДЗ-9 — отчёт для сдачи (Lite)")
    meta_lines(
        doc,
        [
            ("Студент:", "Игорь Кашинцев"),
            ("Курс:", "VibeCoder, Тема 9 — локальная CRM, Supabase, RLS"),
            ("Проект:", "DesignStudio CRM"),
            ("GitHub:", "https://github.com/treshkash323-alt/designstudio-crm"),
            ("Папка:", "Projects/ДЗ-9/designstudio-crm/"),
            ("Запуск:", "localhost:3000 · Studio :54323"),
            ("Дата:", "08.06.2026 · готово к сдаче"),
        ],
    )

    doc.add_heading("1. Цель работы", level=1)
    add_numbered(
        doc,
        [
            "Локальная CRM через Docker + Supabase CLI (без облака).",
            "Два пользователя: менеджер и администратор.",
            "RLS в PostgreSQL: менеджер — только assigned_to = его ID.",
            "Next.js 14: login, /dashboard, /admin/clients.",
            "Проверка: 2 vs 5 лидов — разграничение работает.",
            "Код и документация на GitHub без секретов.",
        ],
    )

    doc.add_heading("2. Отличие от методички", level=1)
    add_table(
        doc,
        ["В методичке", "В реализации"],
        [
            ("trailcamp-crm", "designstudio-crm (Lite)"),
            ("@trailcamp.ru", "@designstudio.ru"),
            ("/admin/leads", "/admin/clients"),
        ],
    )

    doc.add_heading("3. Стек", level=1)
    doc.add_paragraph(
        "Docker · Supabase CLI · PostgreSQL RLS · Next.js 14 · TypeScript · "
        "Tailwind · @supabase/ssr"
    )

    doc.add_heading("4. Безопасность", level=1)
    add_table(
        doc,
        ["Проверка", "Результат"],
        [
            (".env.local в git", "Нет — только .env.local.example"),
            ("RLS на leads, profiles", "Да — manager 2 / admin 5"),
            ("Middleware /admin", "Да — manager редиректится"),
            ("service_role в браузере", "Нет — admin.ts не подключён"),
            ("Security Advisor", "0 errors, 2 warnings (handle_new_user)"),
            ("Доступ из интернета", "Нет — localhost"),
        ],
    )
    doc.add_paragraph(
        "Вывод: утечек секретов в GitHub нет; модель доступа работает. "
        "Учебный локальный проект, не hardened production. Подробно: SECURITY.md."
    )

    doc.add_heading("5. Скриншоты (вставить в документ)", level=1)
    add_table(
        doc,
        ["№", "Что снять", "Подпись"],
        [
            ("1", "Менеджер /dashboard — 2 клиента, RLS (2)", "Рис. 1. Менеджер"),
            ("2", "Админ /admin/clients — 5 клиентов, АДМИН", "Рис. 2. Админ"),
            ("3", "Docker — контейнеры designstudio-crm", "Рис. 3. Docker"),
            ("4", "Studio — таблица leads", "Рис. 4. База"),
            ("5", "Терминал npm run dev", "Рис. 5. Запуск"),
            ("6", "(опц.) Security Advisor — 0 errors", "Рис. 6. Advisor"),
        ],
    )
    doc.add_paragraph(
        "Тестовые аккаунты: manager@designstudio.ru · admin@designstudio.ru. "
        "На скринах замазать anon/service keys."
    )

    doc.add_heading("6. Сценарий проверки", level=1)
    add_numbered(
        doc,
        [
            "Docker → npx supabase start → npm run dev",
            "Регистрация manager@ и admin@designstudio.ru",
            "SQL seed_after_signup.sql в Studio",
            "Admin: выход/вход → /admin/clients — 5 лидов",
            "Manager: /dashboard — 2 лида",
            "Manager → /admin/clients — редирект на /dashboard",
        ],
    )

    doc.add_heading("7. Результат", level=1)
    add_table(
        doc,
        ["Пункт", "Статус"],
        [
            ("Supabase локально, RLS, middleware", "✅"),
            ("Admin + 2 лида менеджеру", "✅"),
            ("Проверка 2 vs 5 лидов", "✅"),
            ("GitHub без секретов", "✅"),
            ("SECURITY.md, TODO_ROADMAP.md", "✅"),
            ("Скрины 1–5 в Word", "☐ вручную"),
        ],
    )

    doc.add_heading("8. Ссылки", level=1)
    add_table(
        doc,
        ["Ресурс", "URL"],
        [
            ("GitHub", "https://github.com/treshkash323-alt/designstudio-crm"),
            ("CRM (локально)", "http://localhost:3000"),
            ("Supabase Studio", "http://localhost:54323"),
        ],
    )

    doc.add_heading("9. Комментарий для преподавателя", level=1)
    doc.add_paragraph(
        "Lite выполнено локально: Docker + Supabase + Next.js, RLS + middleware. "
        "GitHub — публичный репозиторий без .env.local. Аудит безопасности (SECURITY.md). "
        "Advisor: 0 errors. План развития — EDU (admin/student/teacher) и Tilda — "
        "в TODO_ROADMAP.md, после сдачи."
    )

    doc.add_heading("10. Развитие после ДЗ", level=1)
    doc.add_paragraph(
        "См. TODO_ROADMAP.md: образовательная платформа AIKIVAVIORA, роли пользователей, "
        "интеграция с Tilda — отдельный трек, не входит в Lite."
    )

    out = ROOT / "DZ-9_отчёт_для_сдачи.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print("OK guide:", build())
    print("OK report:", build_report())
